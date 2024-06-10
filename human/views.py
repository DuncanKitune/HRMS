from django.shortcuts import render

# Create your views here.
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.models import User
from django.core.mail import send_mail
from django.utils import timezone
from datetime import timedelta
from django.contrib.sites.shortcuts import get_current_site
from django.template.loader import render_to_string, get_template
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_text
from django.contrib.auth.tokens import default_token_generator
from .models import Employee, Payslip, Department, JobGroup, Benefit, Letter, StatutoryDeduction
from .forms import EmployeeCreationForm, IssueLetterForm, Letter, LeaveApplication, LeaveApplicationForm, ContractRenewalForm, StatutoryDeductionFormSet
from xhtml2pdf import pisa
from docx import Document
from openpyxl import Workbook
from django.utils.html import strip_tags
from django.conf import settings
#added on 5/6/24
from django.contrib.auth.decorators import permission_required
from django.http import HttpResponse
from django.template.loader import render_to_string, get_template


@login_required
def admin_dashboard(request):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    employees = Employee.objects.all()
    leave_applications = LeaveApplicationForm.objects.all()
    contract_renewals = ContractRenewalForm.objects.all()
    issued_letters = Letter.objects.all()
    payslips = Payslip.objects.all()
    statutory_deductions = StatutoryDeduction.objects.all()
    return render(request, 'dashboard/admin_dashboard.html', {
        'employees': employees,
        'leave_applications': leave_applications,
        'contract_renewals': contract_renewals,
        'issued_letters': issued_letters,
        'payslips': payslips,
        'statutory_deductions': statutory_deductions,
    })


@login_required
def create_employee(request):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    if request.method == 'POST':
        form = EmployeeCreationForm(request.POST)
        if form.is_valid():
            full_name = form.cleaned_data['full_name']
            email = form.cleaned_data['email']
            department = Department.objects.get(id=department_id)
            department_id = form.cleaned_data['department']
            job_group = form.cleaned_data['job_group']
            salary = form.cleaned_data['salary']
            benefits = form.cleaned_data['benefits']
            user = User.objects.create_user(username=email, email=email, password='defaultpassword')
            employee = Employee.objects.create(
                user=user,
                full_name=full_name,
                department=department,
                salary=salary,
                
                # Fill in other fields here
            )
            employee.benefits.set(benefits)
            employee.save()
            # Send email with activation link
            current_site = get_current_site(request)
            mail_subject = 'Activate your account'
            message = render_to_string('account_activation_email.html', {
                'user': user,
                'domain': current_site.domain,
                'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                'token': default_token_generator.make_token(user),
            })  
           
            # Send email to the new employee with activation link
            send_mail(
                mail_subject,'Welcome to the Company',
                message,f'Hello {full_name},\n\nYou have been added to the company system. Please activate your account and change your default password.\n\nThank you.',
                'admin@example.com',
                [email],
                fail_silently=False,
            )
            return redirect('admin_dashboard')
    else:
        form = EmployeeCreationForm()
    departments = Department.objects.all()
    return render(request, 'dashboard/create_employee.html', {'form': form, 'departments': departments})

def notify_contract_renewal(employee):
    send_mail(
        'Contract Renewal Notification',
        'Your contract is due for renewal in 45 days.',
        'admin@example.com',
        [employee.user.email],
        fail_silently=False,
    )



@login_required
def approve_leave(request, leave_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    leave_application = get_object_or_404(LeaveApplication, id=leave_id)
    if request.method == 'POST':
        leave_application.status = 'approved'
        leave_application.save()
        # Send email notification to employee
        send_mail(
            'Leave Application Approved',
            f'Your leave application from {leave_application.start_date} to {leave_application.end_date} has been approved.',
            'admin@example.com',
            [leave_application.employee.user.email],
            fail_silently=False,
        )
        return redirect('admin_dashboard')
    return render(request, 'dashboard/approve_leave.html', {'leave_application': leave_application})

@login_required
def transfer_employee(request, employee_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    employee = get_object_or_404(Employee, id=employee_id)
    if request.method == 'POST':
        new_department_id = request.POST.get('department')
        new_department = Department.objects.get(id=new_department_id)
        employee.department = new_department
        employee.save()
        return redirect('admin_dashboard')
    departments = Department.objects.all()
    return render(request, 'dashboard/transfer_employee.html', {'employee': employee, 'departments': departments})

def send_email_to_employee(letter, employee):
    subject = f'Issuance of {letter.type}'
    message = f'Dear {employee.full_name},\n\nPlease find attached your {letter.type}.\n\nThank you.'
    recipient_email = employee.user.email
    sender_email = 'admin@example.com'
    send_mail(subject, message, sender_email, [recipient_email], fail_silently=False)


@login_required
def issue_letter(request, employee_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        form = IssueLetterForm(request.POST)
        if form.is_valid():
            letter_type = form.cleaned_data['letter_type']
            subject = form.cleaned_data['subject']
            content = form.cleaned_data['content']
            # Create the letter
            letter = Letter.objects.create(type=letter_type, employee=employee, subject=subject, content=content)
            # Implement logic for issuing the letter (e.g., sending email)
            send_mail(
                subject,
                content,
                'admin@example.com',
                [employee.user.email],
                fail_silently=False,
            )
            # For demonstration purposes, let's assume we're printing the letter type
            print(f"Issuing {letter_type} to {employee.full_name}")
            return redirect('admin_dashboard')
    else:
        form = IssueLetterForm()
    
    return render(request, 'dashboard/issue_letter.html', {'employee': employee, 'form': form})

@login_required
def promote_employee(request, employee_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    employee = get_object_or_404(Employee, id=employee_id)
    if request.method == 'POST':
        new_job_group_id = request.POST.get('job_group')
        new_job_group = JobGroup.objects.get(id=new_job_group_id)
        employee.job_group = new_job_group
        employee.save()
        return redirect('admin_dashboard')
    job_groups = JobGroup.objects.all()
    return render(request, 'dashboard/promote_employee.html', {'employee': employee, 'job_groups': job_groups})

@login_required
def alter_employee_benefits(request, employee_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    employee = get_object_or_404(Employee, id=employee_id)
    if request.method == 'POST':
        new_salary = request.POST.get('new_salary')
        new_benefits = request.POST.getlist('new_benefits')
        employee.salary = new_salary
        employee.benefits.set(new_benefits)
        employee.save()
        return redirect('admin_dashboard')
    benefits = Benefit.objects.all()
    return render(request, 'dashboard/alter_employee_benefits.html', {'employee': employee, 'benefits': benefits})


@login_required
def approve_changes(request, change_id):
    if not request.user.is_superuser:
        return redirect('employee_dashboard')
    change_request = get_object_or_404(ChangeRequest, id=change_id)
    if request.method == 'POST':
        change_request.status = 'approved'
        change_request.save()
        return redirect('admin_dashboard')
    return render(request, 'dashboard/approve_changes.html', {'change_request': change_request})

@login_required
@permission_required('app.view_statutorydeduction', raise_exception=True)
def manage_statutory_deductions(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        formset = StatutoryDeductionFormSet(request.POST, instance=user)
        if formset.is_valid():
            formset.save()
            return redirect('some_view_name')  # Redirect to a relevant view after saving
    else:
        formset = StatutoryDeductionFormSet(instance=user)
    return render(request, 'dashboard/manage_statutory_deductions.html', {'formset': formset, 'user': user})


# Added the below on 5/6/24
# Generate payslips in Word format
@login_required
@permission_required('app.view_payslip', raise_exception=True)
def generate_payslip_pdf(request, payslip_id):
    payslip = get_object_or_404(Payslip, id=payslip_id)
    template_path = 'payslip.html'
    context = {'payslip': payslip}
    template = get_template(template_path)
    html = template.render(context)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="payslip_{payslip.month.strftime("%B_%Y")}.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response

@login_required
@permission_required('app.send_email', raise_exception=True)
def send_payslip_email_view(request, payslip_id):
    payslip = get_object_or_404(Payslip, id=payslip_id)
    payslip.send_payslip_email()
    return HttpResponse('Payslip sent successfully!')

@login_required
@permission_required('users.view_payslip', raise_exception=True)
def generate_payslips_word(request):
    payslips = Payslip.objects.all()
    doc = Document()
    doc.add_heading('Payslips', level=1)

    for payslip in payslips:
        doc.add_paragraph(f'Employee: {payslip.employee.user.username}')
        doc.add_paragraph(f'Month: {payslip.month}')
        doc.add_paragraph(f'Basic Salary: {payslip.basic_salary}')
        doc.add_paragraph(f'Allowances: {payslip.allowances}')
        doc.add_paragraph(f'Deductions: {payslip.deductions}')
        doc.add_paragraph(f'Net Pay: {payslip.net_pay}')
        doc.add_paragraph('')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="payslips.docx"'
    doc.save(response)
    return response

# Generate payslips in Excel format
@login_required
@permission_required('users.view_payslip', raise_exception=True)
def generate_payslips_excel(request):
    payslips = Payslip.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.append(['Employee', 'Month', 'Basic Salary', 'Allowances', 'Deductions', 'Net Pay'])

    for payslip in payslips:
        ws.append([
            payslip.employee.user.username,
            payslip.month,
            payslip.basic_salary,
            payslip.allowances,
            payslip.deductions,
            payslip.net_pay
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="payslips.xlsx"'
    wb.save(response)
    return response


# Manage statutory deductions view
@login_required
@permission_required('users.change_statutorydeduction', raise_exception=True)
def manage_statutory_deductions(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        formset = StatutoryDeductionFormSet(request.POST, instance=user)
        if formset.is_valid():
            formset.save()
            return redirect('some_view_name')  # Redirect to a relevant view after saving
    else:
        formset = StatutoryDeductionFormSet(instance=user)
    return render(request, 'manage_statutory_deductions.html', {'formset': formset, 'user': user})

@login_required
@permission_required('app.view_statutorydeduction', raise_exception=True)
def generate_statutory_deductions_excel(request):
    statutory_deductions = StatutoryDeduction.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.append(['Deduction', 'Percentage'])

    for deduction in statutory_deductions:
        ws.append([deduction.name, deduction.percentage])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="statutory_deductions.xlsx"'
    wb.save(response)
    return response

# Generate monthly payroll reports in PDF format

@login_required
@permission_required('app.view_payslip', raise_exception=True)
def generate_monthly_payroll_pdf(request, month, year):
    payslips = Payslip.objects.filter(month__month=month, month__year=year)
    template_path = 'monthly_payroll_report.html'
    context = {'payslips': payslips, 'month': month, 'year': year}
    template = get_template(template_path)
    html = template.render(context)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="payroll_{month}_{year}.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response

# Generate monthly payroll reports in Word format
@login_required
def generate_payroll_report_word(request):
    payslips = Payslip.objects.all()
    doc = Document()
    doc.add_heading('Payroll Report', level=1)

    for payslip in payslips:
        doc.add_paragraph(f'Employee: {payslip.employee.user.username}')
        doc.add_paragraph(f'Month: {payslip.month}')
        doc.add_paragraph(f'Basic Salary: {payslip.basic_salary}')
        doc.add_paragraph(f'Allowances: {payslip.allowances}')
        doc.add_paragraph(f'Deductions: {payslip.deductions}')
        doc.add_paragraph(f'Net Pay: {payslip.net_pay}')
        doc.add_paragraph('')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="payroll_report.docx"'
    doc.save(response)
    return response

# Generate monthly payroll reports in Excel format
@login_required
def generate_payroll_report_excel(request):
    payslips = Payslip.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.append(['Employee', 'Month', 'Basic Salary', 'Allowances', 'Deductions', 'Net Pay'])

    for payslip in payslips:
        ws.append([
            payslip.employee.user.username,
            payslip.month,
            payslip.basic_salary,
            payslip.allowances,
            payslip.deductions,
            payslip.net_pay
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="payroll_report.xlsx"'
    wb.save(response)
    return response


